#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Módulo de manipulação de Word para o aplicativo Wall Engenharia.
Responsável por preencher o modelo de proposta com os dados extraídos.
"""

import os
import logging
from typing import Dict, List, Any, Optional
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger('word_writer')

class WordWriter:
    """Classe para manipulação de documentos Word e preenchimento de propostas."""
    
    def __init__(self, template_path: str):
        """
        Inicializa o escritor de Word.
        
        Args:
            template_path: Caminho completo para o modelo de documento Word.
        """
        self.template_path = template_path
        self.document = None
        
        # Verificar se o arquivo existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Modelo de documento Word não encontrado: {template_path}")
        
        # Carregar o documento
        try:
            self.document = docx.Document(template_path)
            logger.info(f"Modelo de documento Word carregado com sucesso: {template_path}")
        except Exception as e:
            logger.error(f"Erro ao carregar modelo de documento Word: {e}")
            raise
    
    def find_paragraph_by_text(self, search_text: str, partial_match: bool = False) -> Optional[int]:
        """
        Procura um parágrafo pelo seu texto.
        
        Args:
            search_text: Texto a ser procurado.
            partial_match: Se True, procura por correspondência parcial.
            
        Returns:
            Índice do parágrafo encontrado ou None se não encontrado.
        """
        for i, para in enumerate(self.document.paragraphs):
            if partial_match and search_text.lower() in para.text.lower():
                return i
            elif not partial_match and search_text.lower() == para.text.lower():
                return i
        
        logger.warning(f"Texto '{search_text}' não encontrado no documento")
        return None
    
    def replace_text_in_paragraph(self, paragraph_index: int, new_text: str) -> bool:
        """
        Substitui o texto em um parágrafo específico.
        
        Args:
            paragraph_index: Índice do parágrafo.
            new_text: Novo texto para substituir.
            
        Returns:
            True se a substituição foi bem-sucedida, False caso contrário.
        """
        if 0 <= paragraph_index < len(self.document.paragraphs):
            paragraph = self.document.paragraphs[paragraph_index]
            
            # Preservar a formatação
            runs = paragraph.runs
            if runs:
                # Limpar o parágrafo
                for run in runs:
                    run.text = ""
                
                # Adicionar o novo texto ao primeiro run
                runs[0].text = new_text
            else:
                paragraph.text = new_text
            
            logger.info(f"Texto substituído no parágrafo {paragraph_index}")
            return True
        else:
            logger.warning(f"Índice de parágrafo inválido: {paragraph_index}")
            return False
    
    def replace_text_in_document(self, old_text: str, new_text: str, partial_match: bool = False) -> int:
        """
        Substitui texto em todo o documento.
        
        Args:
            old_text: Texto a ser substituído.
            new_text: Novo texto.
            partial_match: Se True, procura por correspondência parcial.
            
        Returns:
            Número de substituições realizadas.
        """
        count = 0
        
        # Substituir em parágrafos
        for i, para in enumerate(self.document.paragraphs):
            if partial_match and old_text.lower() in para.text.lower():
                para.text = para.text.replace(old_text, new_text)
                count += 1
            elif not partial_match and old_text.lower() == para.text.lower():
                para.text = new_text
                count += 1
        
        # Substituir em tabelas
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if partial_match and old_text.lower() in para.text.lower():
                            para.text = para.text.replace(old_text, new_text)
                            count += 1
                        elif not partial_match and old_text.lower() == para.text.lower():
                            para.text = new_text
                            count += 1
        
        logger.info(f"Realizadas {count} substituições de texto no documento")
        return count
    
    def find_and_replace_placeholder(self, placeholder: str, new_text: str) -> int:
        """
        Procura por um placeholder e o substitui pelo novo texto.
        
        Args:
            placeholder: Placeholder a ser procurado (ex: "{{CLIENTE}}").
            new_text: Texto para substituir o placeholder.
            
        Returns:
            Número de substituições realizadas.
        """
        return self.replace_text_in_document(placeholder, new_text, partial_match=True)
    
    def fill_proposal_with_data(self, data: Dict[str, Any]) -> bool:
        """
        Preenche a proposta com os dados fornecidos.
        
        Args:
            data: Dicionário com os dados para preencher a proposta.
            
        Returns:
            True se o preenchimento foi bem-sucedido, False caso contrário.
        """
        try:
            # Mapeamento de campos para placeholders
            placeholders = {
                'nome_cliente': ['{{CLIENTE}}', '{{NOME_CLIENTE}}'],
                'nome_contato': ['{{CONTATO}}', '{{NOME_CONTATO}}'],
                'email': ['{{EMAIL}}', '{{E-MAIL}}'],
                'telefone': ['{{TELEFONE}}', '{{TEL}}'],
                'escopo': ['{{ESCOPO}}'],
                'prazo': ['{{PRAZO}}'],
                'custo': ['{{CUSTO}}', '{{VALOR}}', '{{PREÇO}}'],
                'garantias': ['{{GARANTIAS}}'],
                'seguro': ['{{SEGURO}}'],
                'nao_inclusos': ['{{NAO_INCLUSOS}}', '{{NÃO_INCLUSOS}}']
            }
            
            # Substituir placeholders pelos valores correspondentes
            for field, field_placeholders in placeholders.items():
                if field in data and data[field]:
                    value = data[field]
                    
                    # Formatar valor numérico para moeda brasileira
                    if field == 'custo' and isinstance(value, (int, float)):
                        value = f"R$ {value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                    
                    # Converter listas para texto com marcadores
                    if isinstance(value, list):
                        value = "\n• " + "\n• ".join(value)
                    
                    # Tentar substituir cada placeholder possível
                    for placeholder in field_placeholders:
                        self.find_and_replace_placeholder(placeholder, str(value))
            
            logger.info("Proposta preenchida com sucesso")
            return True
        except Exception as e:
            logger.error(f"Erro ao preencher proposta: {e}")
            return False
    
    def add_section_if_not_exists(self, section_title: str, content: str) -> bool:
        """
        Adiciona uma seção ao documento se ela não existir.
        
        Args:
            section_title: Título da seção.
            content: Conteúdo da seção.
            
        Returns:
            True se a seção foi adicionada ou já existia, False em caso de erro.
        """
        try:
            # Verificar se a seção já existe
            section_index = self.find_paragraph_by_text(section_title, partial_match=True)
            
            if section_index is not None:
                # A seção já existe, adicionar conteúdo após o título
                self.document.add_paragraph(content, style='Normal')
                return True
            else:
                # Adicionar nova seção no final do documento
                self.document.add_heading(section_title, level=2)
                self.document.add_paragraph(content, style='Normal')
                return True
        except Exception as e:
            logger.error(f"Erro ao adicionar seção: {e}")
            return False
    
    def save_document(self, output_path: str) -> bool:
        """
        Salva o documento em um novo arquivo.
        
        Args:
            output_path: Caminho completo para salvar o documento.
            
        Returns:
            True se o documento foi salvo com sucesso, False caso contrário.
        """
        try:
            # Garantir que o diretório de saída existe
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            self.document.save(output_path)
            logger.info(f"Documento salvo com sucesso: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}")
            return False


# Função auxiliar para uso direto
def generate_proposal(template_path: str, output_path: str, data: Dict[str, Any]) -> bool:
    """
    Gera uma proposta preenchida com os dados fornecidos.
    
    Args:
        template_path: Caminho para o modelo de documento Word.
        output_path: Caminho para salvar a proposta gerada.
        data: Dicionário com os dados para preencher a proposta.
        
    Returns:
        True se a proposta foi gerada com sucesso, False caso contrário.
    """
    try:
        writer = WordWriter(template_path)
        success = writer.fill_proposal_with_data(data)
        
        if success:
            return writer.save_document(output_path)
        else:
            return False
    except Exception as e:
        logger.error(f"Erro ao gerar proposta: {e}")
        return False


if __name__ == "__main__":
    # Teste básico do módulo
    import sys
    import json
    
    if len(sys.argv) > 2:
        template_path = sys.argv[1]
        data_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) > 3 else "proposta_gerada.docx"
    else:
        template_path = input("Digite o caminho para o modelo de documento Word: ")
        data_path = input("Digite o caminho para o arquivo JSON com os dados: ")
        output_path = input("Digite o caminho para salvar a proposta gerada (ou pressione Enter para usar 'proposta_gerada.docx'): ")
        
        if not output_path:
            output_path = "proposta_gerada.docx"
    
    try:
        # Carregar dados do JSON
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Gerar proposta
        success = generate_proposal(template_path, output_path, data)
        
        if success:
            print(f"Proposta gerada com sucesso: {output_path}")
        else:
            print("Erro ao gerar proposta")
            sys.exit(1)
    except Exception as e:
        print(f"Erro: {e}")
        sys.exit(1)
